"""
DOCX to Markdown Converter
Supports mammoth (primary) and python-docx (fallback)
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any

logger = logging.getLogger(__name__)


class DOCXConverter:
    """Convert DOCX documents to Markdown format"""

    def __init__(self, config: Dict[str, Any]):
        """Initialize DOCX converter"""
        self.config = config
        self.strategy = config.get('docx_strategy', 'mammoth')
        self.fallback_order = config.get('fallback_order', {}).get('docx', ['mammoth', 'python-docx'])

    def convert(self, docx_path: str) -> Optional[str]:
        """Convert DOCX to Markdown"""
        try:
            for converter_name in self.fallback_order:
                try:
                    if converter_name == 'mammoth':
                        result = self._convert_with_mammoth(docx_path)
                    elif converter_name == 'python-docx':
                        result = self._convert_with_python_docx(docx_path)
                    else:
                        continue

                    if result:
                        logger.info(f"Successfully converted {docx_path} using {converter_name}")
                        return result

                except ImportError:
                    logger.warning(f"{converter_name} not available")
                    continue
                except Exception as e:
                    logger.warning(f"{converter_name} conversion failed: {e}")
                    continue

            logger.error(f"All converters failed for: {docx_path}")
            return None

        except Exception as e:
            logger.error(f"Error converting DOCX {docx_path}: {e}")
            return None

    def _convert_with_mammoth(self, docx_path: str) -> Optional[str]:
        """Convert DOCX using mammoth (90-95% fidelity to Markdown)"""
        try:
            import mammoth

            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                return result.value

        except ImportError:
            raise
        except Exception as e:
            logger.warning(f"Mammoth conversion failed: {e}")
            return None

    def _convert_with_python_docx(self, docx_path: str) -> Optional[str]:
        """Convert DOCX using python-docx (basic text extraction)"""
        try:
            from docx import Document

            doc = Document(docx_path)
            markdown_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Basic heading detection
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        if level.isdigit():
                            markdown_parts.append(f"{'#' * int(level)} {text}\n\n")
                        else:
                            markdown_parts.append(f"{text}\n\n")
                    else:
                        markdown_parts.append(f"{text}\n\n")

            # Extract tables
            for table in doc.tables:
                markdown_parts.append(self._table_to_markdown(table))
                markdown_parts.append("\n\n")

            return ''.join(markdown_parts)

        except ImportError:
            raise
        except Exception as e:
            logger.warning(f"python-docx conversion failed: {e}")
            return None

    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to markdown"""
        markdown_parts = []

        rows = [[cell.text for cell in row.cells] for row in table.rows]

        if rows:
            # Header
            header = rows[0]
            markdown_parts.append('| ' + ' | '.join(header) + ' |\n')
            markdown_parts.append('| ' + ' | '.join(['---'] * len(header)) + ' |\n')

            # Data rows
            for row in rows[1:]:
                markdown_parts.append('| ' + ' | '.join(row) + ' |\n')

        return ''.join(markdown_parts)
