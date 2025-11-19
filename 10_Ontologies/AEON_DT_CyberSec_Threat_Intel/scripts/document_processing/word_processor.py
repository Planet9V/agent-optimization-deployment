#!/usr/bin/env python3
"""
Word Document Processor for Cybersecurity Threat Intelligence
Extracts text, tables, and structure from DOCX files
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
import json
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Install with: pip install python-docx")

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WordProcessor:
    """Process Word documents and extract structured content"""

    def __init__(self):
        """Initialize Word processor"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

    def extract_metadata(self, docx_path: Path) -> Dict:
        """
        Extract metadata from Word document

        Args:
            docx_path: Path to DOCX file

        Returns:
            Dictionary of metadata
        """
        try:
            doc = Document(docx_path)
            core_props = doc.core_properties

            return {
                'filename': docx_path.name,
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
                'extracted_at': datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"Error extracting metadata from {docx_path}: {e}")
            return {
                'filename': docx_path.name,
                'error': str(e)
            }

    def extract_text_with_structure(self, docx_path: Path) -> Dict:
        """
        Extract text preserving heading hierarchy

        Args:
            docx_path: Path to DOCX file

        Returns:
            Dictionary with structured text
        """
        try:
            doc = Document(docx_path)

            structure = {
                'sections': [],
                'current_section': None,
                'current_subsection': None
            }

            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1

                    if level == 1:
                        # New main section
                        structure['sections'].append({
                            'title': para.text,
                            'level': level,
                            'content': [],
                            'subsections': []
                        })
                        structure['current_section'] = structure['sections'][-1]
                        structure['current_subsection'] = None

                    elif level == 2 and structure['current_section']:
                        # New subsection
                        subsection = {
                            'title': para.text,
                            'level': level,
                            'content': []
                        }
                        structure['current_section']['subsections'].append(subsection)
                        structure['current_subsection'] = subsection

                    elif level >= 3 and structure['current_subsection']:
                        # Lower level heading - add to subsection content
                        structure['current_subsection']['content'].append({
                            'type': 'heading',
                            'level': level,
                            'text': para.text
                        })

                else:
                    # Regular paragraph
                    para_data = {
                        'type': 'paragraph',
                        'text': para.text,
                        'style': para.style.name
                    }

                    # Add to appropriate section
                    if structure['current_subsection']:
                        structure['current_subsection']['content'].append(para_data)
                    elif structure['current_section']:
                        structure['current_section']['content'].append(para_data)
                    else:
                        # Content before any heading
                        if not structure['sections']:
                            structure['sections'].append({
                                'title': 'Introduction',
                                'level': 0,
                                'content': [],
                                'subsections': []
                            })
                            structure['current_section'] = structure['sections'][0]
                        structure['current_section']['content'].append(para_data)

            # Remove helper keys
            del structure['current_section']
            del structure['current_subsection']

            return structure

        except Exception as e:
            logger.error(f"Error extracting structured text from {docx_path}: {e}")
            return {'error': str(e)}

    def extract_tables(self, docx_path: Path) -> List[Dict]:
        """
        Extract tables from Word document

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of table dictionaries
        """
        tables = []

        try:
            doc = Document(docx_path)

            for i, table in enumerate(doc.tables, 1):
                # Extract table data
                rows_data = []

                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_data)

                # Assume first row is header
                headers = rows_data[0] if rows_data else []
                data_rows = rows_data[1:] if len(rows_data) > 1 else []

                tables.append({
                    'table_number': i,
                    'headers': headers,
                    'rows': data_rows,
                    'num_rows': len(data_rows),
                    'num_columns': len(headers)
                })

        except Exception as e:
            logger.error(f"Error extracting tables from {docx_path}: {e}")

        return tables

    def extract_comments(self, docx_path: Path) -> List[Dict]:
        """
        Extract comments from Word document

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of comment dictionaries
        """
        comments = []

        try:
            doc = Document(docx_path)

            # Access document's XML to get comments
            if hasattr(doc, 'part') and hasattr(doc.part, 'package'):
                try:
                    comments_part = doc.part.package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
                    comments_xml = comments_part.blob

                    # Parse comments XML (simplified - full implementation would use lxml)
                    # This is a placeholder for actual XML parsing
                    logger.info("Comments found but detailed extraction requires additional XML parsing")

                except Exception:
                    logger.debug("No comments found in document")

        except Exception as e:
            logger.error(f"Error extracting comments from {docx_path}: {e}")

        return comments

    def extract_styles_and_formatting(self, docx_path: Path) -> Dict:
        """
        Extract style and formatting information

        Args:
            docx_path: Path to DOCX file

        Returns:
            Dictionary of style information
        """
        try:
            doc = Document(docx_path)

            styles_used = set()
            formatting = {
                'bold_count': 0,
                'italic_count': 0,
                'underline_count': 0
            }

            for para in doc.paragraphs:
                styles_used.add(para.style.name)

                # Count formatting
                for run in para.runs:
                    if run.bold:
                        formatting['bold_count'] += 1
                    if run.italic:
                        formatting['italic_count'] += 1
                    if run.underline:
                        formatting['underline_count'] += 1

            return {
                'styles_used': sorted(list(styles_used)),
                'formatting_stats': formatting
            }

        except Exception as e:
            logger.error(f"Error extracting styles from {docx_path}: {e}")
            return {'error': str(e)}

    def process_document(self, docx_path: Path, extract_tables: bool = True,
                        extract_comments: bool = False) -> Dict:
        """
        Comprehensive Word document processing

        Args:
            docx_path: Path to DOCX file
            extract_tables: Extract tables from document
            extract_comments: Extract comments and tracked changes

        Returns:
            Dictionary containing all extracted information
        """
        logger.info(f"Processing Word document: {docx_path}")

        result = {
            'source_file': str(docx_path),
            'metadata': self.extract_metadata(docx_path),
            'structure': self.extract_text_with_structure(docx_path),
            'styles': self.extract_styles_and_formatting(docx_path)
        }

        if extract_tables:
            result['tables'] = self.extract_tables(docx_path)
            logger.info(f"Extracted {len(result['tables'])} tables")

        if extract_comments:
            result['comments'] = self.extract_comments(docx_path)
            logger.info(f"Extracted {len(result['comments'])} comments")

        return result

    def batch_process(self, docx_directory: Path, output_dir: Path) -> List[Dict]:
        """
        Process multiple Word documents in batch

        Args:
            docx_directory: Directory containing DOCX files
            output_dir: Output directory for results

        Returns:
            List of processing results
        """
        output_dir.mkdir(parents=True, exist_ok=True)

        docx_files = list(docx_directory.glob('*.docx'))
        results = []

        for i, docx_file in enumerate(docx_files, 1):
            # Skip temporary files
            if docx_file.name.startswith('~$'):
                continue

            logger.info(f"Processing {i}/{len(docx_files)}: {docx_file.name}")

            try:
                result = self.process_document(docx_file)

                # Save individual result
                output_file = output_dir / f"{docx_file.stem}.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2, ensure_ascii=False)

                results.append({
                    'file': docx_file.name,
                    'status': 'success',
                    'output': str(output_file)
                })

            except Exception as e:
                logger.error(f"Error processing {docx_file}: {e}")
                results.append({
                    'file': docx_file.name,
                    'status': 'error',
                    'error': str(e)
                })

        return results


def main():
    """CLI interface for Word document processing"""
    parser = argparse.ArgumentParser(description='Process Word documents for threat intelligence')
    parser.add_argument('input', help='Input DOCX file or directory')
    parser.add_argument('-o', '--output', help='Output directory', default='tmp')
    parser.add_argument('--no-tables', action='store_true', help='Skip table extraction')
    parser.add_argument('--comments', action='store_true', help='Extract comments and tracked changes')

    args = parser.parse_args()

    # Initialize processor
    processor = WordProcessor()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if input_path.is_file():
        # Process single document
        result = processor.process_document(
            input_path,
            extract_tables=not args.no_tables,
            extract_comments=args.comments
        )

        # Save result
        output_path.mkdir(parents=True, exist_ok=True)
        output_file = output_path / f"{input_path.stem}.json"

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        logger.info(f"Results saved to {output_file}")
        logger.info(f"Sections: {len(result.get('structure', {}).get('sections', []))}")
        logger.info(f"Tables: {len(result.get('tables', []))}")

    elif input_path.is_dir():
        # Batch process directory
        results = processor.batch_process(input_path, output_path)

        # Save summary
        summary_file = output_path / 'processing_summary.json'
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)

        success_count = sum(1 for r in results if r['status'] == 'success')
        logger.info(f"Processed {success_count}/{len(results)} documents successfully")
        logger.info(f"Summary saved to {summary_file}")

    else:
        logger.error(f"Invalid input path: {input_path}")


if __name__ == '__main__':
    main()
