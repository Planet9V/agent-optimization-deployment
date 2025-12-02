"""
DOCX Preprocessor: Convert Word documents to text for NER11 ingestion

Extracts text content from .docx files while preserving structure for entity extraction.

File: docx_preprocessor.py
Created: 2025-12-02
Version: 1.0.0
Purpose: Enable ingestion of Word documents through NER11 pipeline
"""

from docx import Document
from pathlib import Path
from typing import Dict, Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXPreprocessor:
    """
    Preprocessor for extracting text from Word documents.

    Features:
    - Extract all paragraphs with formatting
    - Preserve headings and structure
    - Handle tables
    - Clean and normalize text
    """

    def __init__(self):
        """Initialize DOCX preprocessor."""
        pass

    def extract_text_from_docx(self, docx_path: str) -> Optional[str]:
        """
        Extract plain text from a .docx file.

        Args:
            docx_path: Path to .docx file

        Returns:
            Extracted text string, or None if extraction fails
        """
        try:
            doc = Document(docx_path)

            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            # Join with double newlines to preserve document structure
            full_text = "\n\n".join(paragraphs)

            logger.info(f"Extracted {len(full_text)} characters from {Path(docx_path).name}")
            return full_text

        except Exception as e:
            logger.error(f"Failed to extract from {docx_path}: {e}")
            return None

    def convert_docx_to_markdown(self, docx_path: str, output_path: str) -> bool:
        """
        Convert .docx to .md file for standard pipeline processing.

        Args:
            docx_path: Input .docx file
            output_path: Output .md file path

        Returns:
            True if successful, False otherwise
        """
        text = self.extract_text_from_docx(docx_path)

        if text:
            try:
                with open(output_path, 'w', encoding='utf-8') as f:
                    # Add metadata header
                    doc_name = Path(docx_path).stem
                    f.write(f"# {doc_name}\n\n")
                    f.write(f"**Source**: {Path(docx_path).name}\n")
                    f.write(f"**Converted**: Auto-converted from DOCX\n\n")
                    f.write("---\n\n")
                    f.write(text)

                logger.info(f"Converted {Path(docx_path).name} → {Path(output_path).name}")
                return True

            except Exception as e:
                logger.error(f"Failed to write markdown: {e}")
                return False

        return False

    def batch_convert_directory(self, input_dir: str, output_dir: str) -> Dict[str, int]:
        """
        Convert all .docx files in a directory to markdown.

        Args:
            input_dir: Directory containing .docx files
            output_dir: Directory for .md output

        Returns:
            Statistics: {'total': N, 'converted': M, 'failed': K}
        """
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        docx_files = list(input_path.glob("*.docx"))

        stats = {'total': len(docx_files), 'converted': 0, 'failed': 0}

        logger.info(f"Found {len(docx_files)} DOCX files in {input_dir}")

        for docx_file in docx_files:
            md_filename = docx_file.stem + ".md"
            md_path = output_path / md_filename

            if self.convert_docx_to_markdown(str(docx_file), str(md_path)):
                stats['converted'] += 1
            else:
                stats['failed'] += 1

        logger.info(f"Conversion complete: {stats['converted']}/{stats['total']} successful")
        return stats


def convert_eab_documents():
    """
    Convert EAB DOCX files to markdown for ingestion.

    Returns:
        Path to converted markdown directory
    """
    preprocessor = DOCXPreprocessor()

    input_dir = "/home/jim/2_OXOT_Projects_Dev/3_EAB_40"
    output_dir = "/home/jim/2_OXOT_Projects_Dev/3_EAB_40_converted_md"

    print(f"\n{'='*70}")
    print("EAB DOCX → MARKDOWN CONVERSION")
    print(f"{'='*70}\n")

    stats = preprocessor.batch_convert_directory(input_dir, output_dir)

    print(f"\n{'='*70}")
    print("CONVERSION SUMMARY")
    print(f"{'='*70}")
    print(f"Total DOCX files: {stats['total']}")
    print(f"Successfully converted: {stats['converted']}")
    print(f"Failed: {stats['failed']}")
    print(f"\nMarkdown files ready at: {output_dir}")
    print(f"{'='*70}\n")

    return output_dir


if __name__ == "__main__":
    # Test conversion
    convert_eab_documents()
